#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX转Markdown工具
支持将Word文档(.docx)转换为Markdown格式(.md)

使用方法:
1. 转换单个文件: python docx_to_md.py input.docx
2. 转换并指定输出文件: python docx_to_md.py input.docx output.md
3. 批量转换目录下所有docx文件: python docx_to_md.py /path/to/directory

依赖库:
- python-docx: pip install python-docx
- markdownify: pip install markdownify
"""

import os
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 缺少python-docx库，请运行: pip install python-docx")
    sys.exit(1)

try:
    from markdownify import markdownify as md
except ImportError:
    print("错误: 缺少markdownify库，请运行: pip install markdownify")
    sys.exit(1)

def docx_to_markdown(docx_path, md_path=None):
    """
    将DOCX文件转换为Markdown格式
    
    Args:
        docx_path (str): 输入的DOCX文件路径
        md_path (str, optional): 输出的MD文件路径，如果不指定则自动生成
    
    Returns:
        str: 输出文件路径
    """
    try:
        # 读取DOCX文件
        doc = Document(docx_path)
        
        # 提取文本内容
        content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # 根据段落样式添加Markdown格式
                style_name = ""
                if paragraph.style and paragraph.style.name:
                    style_name = paragraph.style.name.lower()
                
                if 'heading 1' in style_name or 'title' in style_name:
                    content.append(f"# {text}")
                elif 'heading 2' in style_name:
                    content.append(f"## {text}")
                elif 'heading 3' in style_name:
                    content.append(f"### {text}")
                elif 'heading 4' in style_name:
                    content.append(f"#### {text}")
                elif 'heading 5' in style_name:
                    content.append(f"##### {text}")
                elif 'heading 6' in style_name:
                    content.append(f"###### {text}")
                else:
                    content.append(text)
                
                content.append("")  # 添加空行
        
        # 处理表格
        for table in doc.tables:
            table_md = []
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                
                table_md.append("| " + " | ".join(row_data) + " |")
                
                # 添加表头分隔符
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(row_data)) + " |"
                    table_md.append(separator)
            
            if table_md:
                content.extend(table_md)
                content.append("")  # 添加空行
        
        # 生成输出文件路径
        if md_path is None:
            docx_file = Path(docx_path)
            md_path = docx_file.parent / f"{docx_file.stem}.md"
        
        # 写入Markdown文件
        markdown_content = "\n".join(content)
        
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"转换成功: {docx_path} -> {md_path}")
        return str(md_path)
        
    except Exception as e:
        print(f"转换失败 {docx_path}: {str(e)}")
        return None

def batch_convert(directory_path):
    """
    批量转换目录下的所有DOCX文件
    
    Args:
        directory_path (str): 目录路径
    """
    directory = Path(directory_path)
    
    if not directory.exists() or not directory.is_dir():
        print(f"错误: 目录不存在 - {directory_path}")
        return
    
    docx_files = list(directory.glob("*.docx")) + list(directory.glob("*.xml"))
    
    if not docx_files:
        print(f"在目录 {directory_path} 中未找到DOCX或XML文件")
        return
    
    print(f"找到 {len(docx_files)} 个文件，开始批量转换...")
    
    success_count = 0
    for docx_file in docx_files:
        if docx_to_markdown(str(docx_file)):
            success_count += 1
    
    print(f"批量转换完成: {success_count}/{len(docx_files)} 个文件转换成功")

def main():
    parser = argparse.ArgumentParser(description='DOCX转Markdown工具')
    parser.add_argument('input', help='输入的DOCX文件路径或目录路径')
    parser.add_argument('output', nargs='?', help='输出的MD文件路径（可选）')
    parser.add_argument('--batch', '-b', action='store_true', help='批量转换模式')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    
    if not input_path.exists():
        print(f"错误: 输入路径不存在 - {args.input}")
        sys.exit(1)
    
    if input_path.is_dir() or args.batch:
        # 批量转换模式
        batch_convert(str(input_path))
    elif input_path.suffix.lower() in ['.docx', '.xml']:
        # 单文件转换模式
        docx_to_markdown(str(input_path), args.output)
    else:
        print("错误: 输入文件必须是.docx或.xml格式")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) == 1:
        print("DOCX转Markdown工具")
        print("\n使用方法:")
        print("1. 转换单个文件: python docx_to_md.py input.docx")
        print("2. 转换并指定输出: python docx_to_md.py input.docx output.md")
        print("3. 批量转换目录: python docx_to_md.py /path/to/directory")
        print("\n依赖库安装:")
        print("pip install python-docx markdownify")
    else:
        main()